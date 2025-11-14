import argparse
import sys
from copy import deepcopy
from typing import List, Optional, Dict, Tuple
import os

try:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
except Exception as e:  # pragma: no cover
    print("Missing dependency: python-docx. Install with: pip install python-docx", file=sys.stderr)
    raise

# Optional, for robust single-file merging
try:
    from docxcompose.composer import Composer  # type: ignore
    HAS_COMPOSER = True
except Exception:
    HAS_COMPOSER = False

try:
    import openpyxl
except Exception as e:  # pragma: no cover
    print("Missing dependency: openpyxl. Install with: pip install openpyxl", file=sys.stderr)
    raise


"""
# ---------------------------
# Excel helpers
# ---------------------------
"""

def read_names_from_excel(
    path: str,
    sheet: Optional[str] = None,
) -> List[str]:
    """
    Read nominative full names from the first column (A) of the sheet.
    Reads from the FIRST row as data (no header assumed).
    """
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb[sheet] if sheet else wb[wb.sheetnames[0]]
    names: List[str] = []
    for row in ws.iter_rows(min_row=1, max_col=1, values_only=True):
        val = row[0]
        if val and str(val).strip():
            names.append(str(val).strip())
    return names


# ---------------------------
# Ukrainian dative heuristics (fallback)
# ---------------------------

VOWELS = set("аеєиіїоуюя")


def _endswith_any(s: str, suffixes: Tuple[str, ...]) -> Optional[str]:
    for suf in suffixes:
        if s.endswith(suf):
            return suf
    return None


def _is_vowel(ch: str) -> bool:
    return ch.lower() in VOWELS


def dative_first_name(name: str) -> str:
    n = name.strip()
    low = n.lower()
    if not n:
        return n
    if low.endswith("а"):
        return n[:-1] + ("і" if n[-1].islower() else "І")
    if low.endswith("я"):
        return n[:-1] + ("ї" if n[-1].islower() else "Ї")
    if low.endswith("й"):
        return n[:-1] + ("ю" if n[-1].islower() else "Ю")
    if low.endswith("о"):
        base = n[:-1]
        add = "ові"
        # Capitalization heuristic
        return base + (add if n[-1].islower() else add.capitalize())
    # Consonant or soft sign
    if not _is_vowel(low[-1]) or low.endswith("ь"):
        add = "ові"
        return n + (add if n[-1].islower() else add.capitalize())
    return n  # fallback


def dative_patronymic(p: str, gender: Optional[str] = None) -> str:
    if not p:
        return ""
    low = p.lower()
    male_suf = ("ович", "йович", "евич", "льович", "євич")
    fem_suf = ("івна", "ївна", "евна", "ївна", "євна")
    if _endswith_any(low, male_suf):
        return p + ("у" if p[-1].islower() else "У")
    if _endswith_any(low, fem_suf):
        # a -> i
        return p[:-1] + ("і" if p[-1].islower() else "І")
    # Heuristic by gender
    if gender == "f":
        if low.endswith("а"):
            return p[:-1] + ("і" if p[-1].islower() else "І")
        if low.endswith("я"):
            return p[:-1] + ("ї" if p[-1].islower() else "Ї")
        return p
    # default masculine-like
    if not _is_vowel(low[-1]) or low.endswith("ь") or low.endswith("й"):
        add = "у"
        return p + (add if p[-1].islower() else add.upper())
    return p


def dative_surname(surname: str, gender: Optional[str] = None) -> str:
    s = surname.strip()
    low = s.lower()
    if not s:
        return s
    if low.endswith("енко"):
        return s[:-1] + ("у" if s[-1].islower() else "У")
    if low.endswith("ко"):
        return s[:-1] + ("у" if s[-1].islower() else "У")
    if low.endswith("о"):
        return s[:-1] + ("у" if s[-1].islower() else "У")
    if low.endswith("а"):
        return s[:-1] + ("і" if s[-1].islower() else "І")
    if low.endswith("я"):
        return s[:-1] + ("ї" if s[-1].islower() else "Ї")
    if low.endswith("й"):
        return s[:-1] + ("ю" if s[-1].islower() else "Ю")
    # Consonant endings
    if not _is_vowel(low[-1]) or low.endswith("ь"):
        add = "у"
        return s + (add if s[-1].islower() else add.upper())
    return s


def guess_gender_from_patronymic(p: Optional[str]) -> Optional[str]:
    if not p:
        return None
    low = p.lower()
    if low.endswith("ович") or low.endswith("йович") or low.endswith("евич") or low.endswith("льович") or low.endswith("євич"):
        return "m"
    if low.endswith("івна") or low.endswith("ївна") or low.endswith("евна") or low.endswith("євна"):
        return "f"
    return None


def split_fullname(fullname: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    parts = [p for p in fullname.strip().split() if p]
    if not parts:
        return None, None, None
    if len(parts) == 1:
        return None, parts[0], None
    if len(parts) == 2:
        return parts[0], parts[1], None
    # 3 or more: assume Surname Name Patronymic, ignore extras
    return parts[0], parts[1], parts[2]


def to_dative_fullname(fullname: str) -> str:
    s, n, p = split_fullname(fullname)
    gender = guess_gender_from_patronymic(p)
    out = []
    if s:
        out.append(dative_surname(s, gender))
    if n:
        out.append(dative_first_name(n))
    if p:
        out.append(dative_patronymic(p, gender))
    print(out, s, n, p)
    return " ".join(out) if out else fullname


# ---------------------------
# DOCX processing
# ---------------------------

def _replace_in_element_textnodes(root, mapping: Dict[str, str]):
    # Replace in a given XML root across all w:t nodes with namespace mapping.
    try:
        ns = {'w': root.nsmap.get('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')}
        for t in root.xpath('.//w:t', namespaces=ns):
            if t.text:
                txt = t.text
                for key, val in mapping.items():
                    if key in txt:
                        txt = txt.replace(key, val)
                t.text = txt
    except Exception:
        pass


def _iter_paragraphs(container):
    # Yield all paragraphs within a container, including those nested in tables.
    for p in getattr(container, 'paragraphs', []):
        yield p
    for tbl in getattr(container, 'tables', []):
        for row in tbl.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _replace_in_paragraph_runs(paragraph, mapping: Dict[str, str]):
    runs = getattr(paragraph, 'runs', None)
    if not runs:
        return
    original = ''.join(r.text or '' for r in runs)
    replaced = original
    for k, v in mapping.items():
        if k:
            replaced = replaced.replace(k, v)
    if replaced != original:
        runs[0].text = replaced
        for r in runs[1:]:
            r.text = ''


def replace_everywhere(doc: Document, mapping: Dict[str, str]):
    # First, low-level XML replacement (fast for simple, unsplit placeholders)
    _replace_in_element_textnodes(doc._element, mapping)

    # Then robust replacement across runs for all paragraphs in document, headers, and footers
    for p in _iter_paragraphs(doc):
        _replace_in_paragraph_runs(p, mapping)

    for section in doc.sections:
        for container in (section.header, section.footer):
            try:
                _replace_in_element_textnodes(container._element, mapping)
            except Exception:
                pass
            for p in _iter_paragraphs(container):
                _replace_in_paragraph_runs(p, mapping)


def flatten_headers_into_body(doc: Document):
    # Move header/footer contents into the document body and clear headers/footers.
    body_el = doc._element.body
    for section in doc.sections:
        try:
            hdr_el = section.header._element
            for child in list(hdr_el):
                body_el.append(deepcopy(child))
                hdr_el.remove(child)
        except Exception:
            pass
        try:
            ftr_el = section.footer._element
            for child in list(ftr_el):
                # Typically footer not needed; skip moving to body to avoid extra content at bottom
                ftr_el.remove(child)
        except Exception:
            pass


def generate_from_template(
    template_path: str,
    names: List[str],
    placeholders: List[str],
    out_path: str,
    keep_first_page: bool = True,
    flatten_for_single: bool = True,
):
    if not names:
        raise ValueError("No names to process")

    # Build per-person docs from a fresh template, then stitch bodies together.
    out_doc = None

    def build_replacement_map(fullname: str) -> Dict[str, str]:
        repl = to_dative_fullname(fullname)
        return {ph: repl for ph in placeholders}

    def _clone_header_footer(src_doc: Document, dst_doc: Document, dst_section):
        try:
            dst_section.header.is_linked_to_previous = False
            # Clear dst header content
            dst_hdr_el = dst_section.header._element
            for el in list(dst_hdr_el):
                dst_hdr_el.remove(el)
            # Copy src header content
            src_hdr_el = src_doc.sections[0].header._element
            for child in list(src_hdr_el):
                dst_hdr_el.append(deepcopy(child))
        except Exception:
            pass
        try:
            dst_section.footer.is_linked_to_previous = False
            dst_ftr_el = dst_section.footer._element
            for el in list(dst_ftr_el):
                dst_ftr_el.remove(el)
            src_ftr_el = src_doc.sections[0].footer._element
            for child in list(src_ftr_el):
                dst_ftr_el.append(deepcopy(child))
        except Exception:
            pass

    def _copy_section_settings(src_section, dst_section):
        try:
            src_sp = src_section._sectPr
            dst_sp = dst_section._sectPr
            # Remove existing non-header/footer settings in destination
            for el in list(dst_sp):
                ln = el.tag.split('}')[-1]
                if ln in ("headerReference", "footerReference"):
                    continue
                dst_sp.remove(el)
            # Copy over settings except header/footer refs and type
            for el in list(src_sp):
                ln = el.tag.split('}')[-1]
                if ln in ("headerReference", "footerReference", "type"):
                    continue
                dst_sp.append(deepcopy(el))
        except Exception:
            pass

    for idx, person in enumerate(names):
        doc_i = Document(template_path)
        replace_everywhere(doc_i, build_replacement_map(person))

        if flatten_for_single:
            flatten_headers_into_body(doc_i)

        if out_doc is None:
            out_doc = doc_i
        else:
            new_sec = out_doc.add_section(WD_SECTION_START.NEW_PAGE)
            # Preserve page setup from the template page
            try:
                _copy_section_settings(doc_i.sections[0], new_sec)
            except Exception:
                pass
            for child in list(doc_i._element.body):
                ln = child.tag.split('}')[-1]
                if ln == 'sectPr':
                    continue
                out_doc._element.body.append(deepcopy(child))

    if out_doc is None:
        raise ValueError("Failed to build output document")

    out_doc.save(out_path)


def generate_separate_files(
    template_path: str,
    names: List[str],
    placeholders: List[str],
    out_dir: str = "out",
):
    os.makedirs(out_dir, exist_ok=True)
    for i, person in enumerate(names, 1):
        doc = Document(template_path)
        repl = to_dative_fullname(person)
        replace_everywhere(doc, {ph: repl for ph in placeholders})
        # Safe filename: use index and last name if present
        s, n, p = split_fullname(person)
        label = s or (n or f"{i}")
        fname = f"{i:03d}_{label}.docx"
        out_path = os.path.join(out_dir, fname)
        doc.save(out_path)


def merge_docx_in_dir(in_dir: str, out_path: str):
    files = sorted(
        [os.path.join(in_dir, f) for f in os.listdir(in_dir) if f.lower().endswith('.docx')]
    )
    if not files:
        raise ValueError(f"No .docx files found in directory: {in_dir}")
    if HAS_COMPOSER:
        master = Document(files[0])
        composer = Composer(master)
        for f in files[1:]:
            composer.append(Document(f))
        composer.save(out_path)
    else:
        raise RuntimeError(
            "docxcompose is not installed. Install with: pip install docxcompose"
        )


# Removed complex header/column detection; names are always in first column


def main(argv: Optional[List[str]] = None) -> int:
    ap = argparse.ArgumentParser(description="Generate medals DOCX by duplicating a template and replacing a name placeholder")
    ap.add_argument("--template", default="medals.docx", help="Path to DOCX template (default: medals.docx)")
    ap.add_argument("--excel", default="listm.xlsx", help="Path to Excel with names (default: listm.xlsx)")
    ap.add_argument("--sheet", default=None, help="Excel sheet name (default: first sheet)")
    # Names are always in the first column; no column/no-header options
    ap.add_argument("--placeholder", action="append", default=None, help="Text to replace; can be repeated. Default targets 'Гурову Денису Сергійовичу' and common variants.")
    ap.add_argument("--output", default="medals_out.docx", help="Output DOCX file (default: medals_out.docx)")
    # Default to separate files for reliability (headers/footers). Allow overriding with --single.
    ap.add_argument("--separate", dest="separate", action="store_true", default=True, help="Generate separate DOCX files per person into ./out directory (default)")
    ap.add_argument("--single", dest="separate", action="store_false", help="Generate one multi-page DOCX instead of separate files")
    # Always replace full PІБ in dative; no replace mode option
    ap.add_argument("--out-dir", default="out", help="Directory to place per-person DOCX before merging (default: out)")
    args = ap.parse_args(argv)

    # Read names (first column, skipping header row)
    names = read_names_from_excel(
        path=args.excel,
        sheet=args.sheet,
    )

    if not names:
        print("No names found in Excel.")
        return 1

    # Resolve placeholders list. Default to typical template texts for your case.
    if args.placeholder:
        placeholders = args.placeholder
    else:
        placeholders = [
            "Гурову Денису Сергійовичу",  # dative full in template
            "Гуров Денис Сергійович",     # nominative full variant
            "Гуров",                      # bare surname variant
        ]

    # Generate document(s)
    if args.separate:
        generate_separate_files(
            template_path=args.template,
            names=names,
            placeholders=placeholders,
            out_dir=args.out_dir,
        )
        print(f"Done. Generated separate files in ./{args.out_dir}")
        print("Note: Converted from nominative using heuristic rules. Please proofread.")
        return 0
    else:
        # Build single file by composing per-person docs via docxcompose.Composer
        # 1) Generate reliable per-person DOCX files into out_dir
        generate_separate_files(
            template_path=args.template,
            names=names,
            placeholders=placeholders,
            out_dir=args.out_dir,
        )
        # 2) Merge them into a single DOCX preserving layout/headers/footers
        try:
            merge_docx_in_dir(args.out_dir, args.output)
        except Exception as e:
            print(str(e))
            print("Merging requires 'docxcompose'. Install with: pip install docxcompose")
            return 1
        print(f"Done. Generated single file (composer): {args.output} (pages: {len(names)})")

    print("Note: Converted from nominative using heuristic rules. Please proofread.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
